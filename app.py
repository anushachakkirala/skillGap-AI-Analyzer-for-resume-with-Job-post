"""
AI Resume-JD Matcher with Enhanced Skill Extraction and Sentence-BERT
Complete integrated application with OTP-based login
"""

import streamlit as st
import PyPDF2
import docx
import re
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from collections import Counter
import numpy as np

from sklearn.metrics.pairwise import cosine_similarity
import io
from datetime import datetime
import base64
import streamlit.components.v1 as components
import os
from PIL import Image

# Import authentication modules
from auth_database import AuthDatabase
from email_service import EmailService
from auth_ui import show_signup_page, show_login_page, show_otp_page, show_profile_page

# Try importing advanced NLP libraries
try:
    import spacy
    SPACY_AVAILABLE = True
    try:
        nlp = spacy.load("en_core_web_sm")
    except:
        SPACY_AVAILABLE = False
        nlp = None
except:
    SPACY_AVAILABLE = False
    nlp = None

try:
    from sentence_transformers import SentenceTransformer
    SBERT_AVAILABLE = True
except:
    SBERT_AVAILABLE = False

try:
    from nltk.stem import PorterStemmer
    from nltk.tokenize import word_tokenize
    import nltk
    # Download required NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    NLTK_AVAILABLE = True
except:
    NLTK_AVAILABLE = False

# ============== SESSION STATE INITIALIZATION ==============

def init_session_state():
    """Initialize all session state variables"""
    # Authentication states
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'auth_page' not in st.session_state:
        st.session_state.auth_page = 'login'
    if 'user' not in st.session_state:
        st.session_state.user = None
    if 'show_profile' not in st.session_state:
        st.session_state.show_profile = False
    
    # Existing app states
    if 'parsed_data' not in st.session_state:
        st.session_state.parsed_data = {}
    if 'extracted_skills' not in st.session_state:
        st.session_state.extracted_skills = {}
    if 'analysis_complete' not in st.session_state:
        st.session_state.analysis_complete = False
    if 'embeddings' not in st.session_state:
        st.session_state.embeddings = {}
    if 'sbert_model' not in st.session_state:
        st.session_state.sbert_model = None

# ============== CATEGORIZED SKILL DATABASES ==============

CATEGORIZED_SKILLS = {
    'programming_languages': {
        'skills': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust', 'swift',
                  'kotlin', 'php', 'r', 'scala', 'perl', 'objective-c', 'dart', 'elixir', 'haskell', 'lua'],
        'category_type': 'technical'
    },
    
    'frontend_frameworks': {
        'skills': ['react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt.js', 'gatsby', 'ember.js',
                  'backbone.js', 'jquery', 'redux', 'mobx', 'recoil', 'zustand'],
        'category_type': 'technical'
    },
    
    'backend_frameworks': {
        'skills': ['node.js', 'express', 'express.js', 'nest.js', 'fastify', 'django', 'flask', 'fastapi',
                  'spring', 'spring boot', 'asp.net', 'asp.net core', 'laravel', 'symfony', 'rails'],
        'category_type': 'technical'
    },
    
    'databases': {
        'skills': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'cassandra', 'dynamodb', 'oracle',
                  'sql server', 'sqlite', 'elasticsearch', 'neo4j', 'firebase', 'supabase'],
        'category_type': 'technical'
    },
    
    'cloud_platforms': {
        'skills': ['aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp', 'google cloud platform',
                  'ec2', 's3', 'lambda', 'kubernetes', 'k8s', 'docker', 'terraform', 'cloudformation'],
        'category_type': 'technical'
    },
    
    'ml_libraries': {
        'skills': ['tensorflow', 'keras', 'pytorch', 'scikit-learn', 'sklearn', 'pandas', 'numpy',
                  'scipy', 'matplotlib', 'seaborn', 'plotly', 'opencv', 'spacy', 'nltk', 'transformers',
                  'hugging face', 'xgboost', 'lightgbm', 'catboost'],
        'category_type': 'technical'
    },
    
    'ml_concepts': {
        'skills': ['machine learning', 'deep learning', 'artificial intelligence', 'nlp',
                  'natural language processing', 'computer vision', 'neural networks', 'cnn', 'rnn',
                  'lstm', 'transformer', 'bert', 'gpt', 'llm', 'generative ai', 'reinforcement learning'],
        'category_type': 'technical'
    },
    
    'data_tools': {
        'skills': ['spark', 'pyspark', 'hadoop', 'airflow', 'dbt', 'tableau', 'power bi', 'looker',
                  'jupyter', 'databricks', 'mlflow', 'sagemaker'],
        'category_type': 'technical'
    },
    
    'devops_tools': {
        'skills': ['jenkins', 'gitlab ci', 'github actions', 'circleci', 'ansible', 'chef', 'puppet',
                  'prometheus', 'grafana', 'datadog', 'elk stack', 'nginx', 'apache'],
        'category_type': 'technical'
    },
    
    'testing_frameworks': {
        'skills': ['jest', 'mocha', 'pytest', 'unittest', 'selenium', 'cypress', 'playwright',
                  'junit', 'testng', 'mockito'],
        'category_type': 'technical'
    },
    
    'mobile_frameworks': {
        'skills': ['react native', 'flutter', 'xamarin', 'ionic', 'swiftui', 'jetpack compose'],
        'category_type': 'technical'
    },
    
    'web_technologies': {
        'skills': ['html', 'html5', 'css', 'css3', 'sass', 'scss', 'tailwind', 'bootstrap',
                  'webpack', 'vite', 'babel'],
        'category_type': 'technical'
    },
    
    'api_protocols': {
        'skills': ['rest api', 'restful', 'graphql', 'grpc', 'soap', 'websocket', 'oauth', 'jwt'],
        'category_type': 'technical'
    },
    
    'version_control': {
        'skills': ['git', 'github', 'gitlab', 'bitbucket', 'svn'],
        'category_type': 'technical'
    },
    
    'soft_skills_leadership': {
        'skills': ['leadership', 'team management', 'mentoring', 'coaching', 'strategic thinking',
                  'decision making', 'delegation'],
        'category_type': 'soft'
    },
    
    'soft_skills_communication': {
        'skills': ['communication', 'presentation', 'negotiation', 'active listening', 'empathy',
                  'interpersonal skills', 'persuasion'],
        'category_type': 'soft'
    },
    
    'soft_skills_technical': {
        'skills': ['problem solving', 'critical thinking', 'analytical', 'attention to detail',
                  'debugging', 'troubleshooting'],
        'category_type': 'soft'
    },
    
    'soft_skills_collaboration': {
        'skills': ['teamwork', 'collaboration', 'cross-functional collaboration', 'conflict resolution',
                  'stakeholder management'],
        'category_type': 'soft'
    },
    
    'soft_skills_personal': {
        'skills': ['adaptability', 'flexibility', 'time management', 'creativity', 'initiative',
                  'self-motivated', 'continuous learning'],
        'category_type': 'soft'
    },
    
    'methodologies': {
        'skills': ['agile', 'scrum', 'kanban', 'devops', 'tdd', 'bdd', 'ci/cd', 'microservices'],
        'category_type': 'technical'
    },
    
    'certifications': {
        'skills': ['aws certified', 'azure certified', 'gcp certified', 'pmp', 'csm', 'cissp',
                  'comptia', 'cka', 'ckad'],
        'category_type': 'certification'
    }
}

# Flatten for quick lookup
ALL_SKILLS_DICT = {}
for category, data in CATEGORIZED_SKILLS.items():
    for skill in data['skills']:
        ALL_SKILLS_DICT[skill] = {
            'category': category,
            'type': data['category_type']
        }

# ============== ENHANCED SKILL EXTRACTION ==============

def extract_noun_chunks(text):
    """Extract noun chunks using spaCy"""
    if not SPACY_AVAILABLE or nlp is None:
        return []
    
    doc = nlp(text[:100000])  # Limit text length for performance
    noun_chunks = [chunk.text.lower().strip() for chunk in doc.noun_chunks]
    return noun_chunks

def extract_entities(text):
    """Extract named entities using spaCy"""
    if not SPACY_AVAILABLE or nlp is None:
        return []
    
    doc = nlp(text[:100000])
    entities = [(ent.text.lower().strip(), ent.label_) for ent in doc.ents]
    return entities

def normalize_text(text, use_stemming=False):
    """Normalize text with optional stemming"""
    text_lower = text.lower()
    
    if use_stemming and NLTK_AVAILABLE:
        stemmer = PorterStemmer()
        words = word_tokenize(text_lower)
        stemmed = [stemmer.stem(word) for word in words]
        return ' '.join(stemmed)
    
    return text_lower

def extract_skills_enhanced(text, confidence_threshold=0.6, use_noun_chunks=True, use_stemming=False):
    """
    Enhanced skill extraction using multiple methods:
    1. Direct pattern matching
    2. Noun chunk extraction (if spaCy available)
    3. Named entity recognition
    4. Context-aware scoring
    """
    
    normalized_text = normalize_text(text, use_stemming)
    sentences = re.split(r'[.!?\n]+', text)
    
    # Method 1: Direct pattern matching
    found_skills = {}
    
    for skill, metadata in ALL_SKILLS_DICT.items():
        pattern = r'\b' + re.escape(skill) + r'\b'
        matches = list(re.finditer(pattern, normalized_text))
        
        if matches:
            evidence = []
            for sentence in sentences:
                if skill in sentence.lower():
                    evidence.append(sentence.strip())
            
            # Calculate confidence based on frequency and context
            frequency = len(matches)
            confidence = min(0.95, 0.6 + (frequency * 0.1))
            
            # Boost confidence if found in section headers
            if any(header in normalized_text for header in ['skills:', 'technologies:', 'expertise:']):
                confidence = min(0.98, confidence + 0.1)
            
            if confidence >= confidence_threshold:
                found_skills[skill] = {
                    'category': metadata['category'],
                    'type': metadata['type'],
                    'count': frequency,
                    'confidence': confidence,
                    'evidence': evidence[:3],
                    'extraction_method': 'pattern_matching'
                }
    
    # Method 2: Noun chunk extraction
    if use_noun_chunks and SPACY_AVAILABLE:
        noun_chunks = extract_noun_chunks(text)
        for chunk in noun_chunks:
            chunk_normalized = normalize_text(chunk, use_stemming)
            if chunk_normalized in ALL_SKILLS_DICT and chunk_normalized not in found_skills:
                metadata = ALL_SKILLS_DICT[chunk_normalized]
                found_skills[chunk_normalized] = {
                    'category': metadata['category'],
                    'type': metadata['type'],
                    'count': 1,
                    'confidence': 0.75,
                    'evidence': [chunk],
                    'extraction_method': 'noun_chunk'
                }
    
    # Method 3: Entity recognition (for certifications and tools)
    if SPACY_AVAILABLE:
        entities = extract_entities(text)
        for entity_text, entity_label in entities:
            entity_normalized = normalize_text(entity_text, use_stemming)
            if entity_normalized in ALL_SKILLS_DICT and entity_normalized not in found_skills:
                metadata = ALL_SKILLS_DICT[entity_normalized]
                found_skills[entity_normalized] = {
                    'category': metadata['category'],
                    'type': metadata['type'],
                    'count': 1,
                    'confidence': 0.80,
                    'evidence': [entity_text],
                    'extraction_method': 'entity_recognition'
                }
    
    # Categorize results
    categorized_results = {}
    for category in CATEGORIZED_SKILLS.keys():
        categorized_results[category] = {}
    
    for skill, data in found_skills.items():
        category = data['category']
        categorized_results[category][skill] = data
    
    return categorized_results

# ============== SENTENCE-BERT EMBEDDINGS ==============

def load_sbert_model(model_name='all-MiniLM-L6-v2'):
    """Load Sentence-BERT model"""
    if not SBERT_AVAILABLE:
        return None
    
    try:
        model = SentenceTransformer(model_name)
        return model
    except Exception as e:
        st.error(f"Failed to load SBERT model: {str(e)}")
        return None

def compute_embeddings(texts, model, normalize=True):
    """Compute embeddings for texts"""
    if model is None:
        return None
    
    try:
        embeddings = model.encode(texts, convert_to_tensor=False, show_progress_bar=False)
        if normalize:
            # Normalize embeddings
            norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
            embeddings = embeddings / norms
        return embeddings
    except Exception as e:
        st.error(f"Embedding computation failed: {str(e)}")
        return None

def compute_similarity_matrix(embeddings1, embeddings2):
    """Compute cosine similarity matrix between two sets of embeddings"""
    if embeddings1 is None or embeddings2 is None:
        return None
    
    # Cosine similarity (already normalized, so just dot product)
    similarity_matrix = np.dot(embeddings1, embeddings2.T)
    return similarity_matrix

# ============== ENHANCED GAP ANALYSIS ==============

def perform_enhanced_gap_analysis(resume_skills, jd_skills, similarity_threshold=0.6, 
                                 embeddings_resume=None, embeddings_jd=None):
    """
    Enhanced gap analysis with:
    - Embedding-based similarity scoring
    - Weighted importance calculation
    - Best matching phrase identification
    """
    
    gaps = []
    
    # Flatten skills for analysis
    jd_skills_flat = {}
    resume_skills_flat = {}
    
    for category, skills in jd_skills.items():
        for skill, data in skills.items():
            jd_skills_flat[skill] = data
    
    for category, skills in resume_skills.items():
        for skill, data in skills.items():
            resume_skills_flat[skill] = data
    
    # Prepare skill lists for embedding
    jd_skill_list = list(jd_skills_flat.keys())
    resume_skill_list = list(resume_skills_flat.keys())
    
    # Compute similarity scores
    if embeddings_jd is not None and embeddings_resume is not None:
        similarity_matrix = compute_similarity_matrix(embeddings_jd, embeddings_resume)
    else:
        similarity_matrix = None
    
    # Analyze each JD skill
    for idx, jd_skill in enumerate(jd_skill_list):
        jd_data = jd_skills_flat[jd_skill]
        
        # Calculate importance (based on frequency in JD)
        importance = min(1.0, 0.5 + (jd_data['count'] * 0.2))
        
        # Find best matching resume skill
        if jd_skill in resume_skills_flat:
            # Exact match
            match_score = 1.0
            best_match = jd_skill
            matched = True
        else:
            # Find best semantic match using embeddings
            if similarity_matrix is not None and idx < len(similarity_matrix):
                best_idx = np.argmax(similarity_matrix[idx])
                match_score = similarity_matrix[idx][best_idx]
                best_match = resume_skill_list[best_idx] if best_idx < len(resume_skill_list) else None
                matched = match_score >= similarity_threshold
            else:
                # Fallback to string similarity
                match_score = 0.0
                best_match = None
                matched = False
        
        # Calculate gap score
        gap_score = importance * (1 - match_score)
        
        gaps.append({
            'jd_skill': jd_skill,
            'category': jd_data['category'],
            'type': jd_data['type'],
            'importance': importance,
            'match_score': match_score,
            'best_resume_match': best_match,
            'matched': matched,
            'gap_score': gap_score,
            'evidence': jd_data.get('evidence', [])
        })
    
    # Sort by gap score (highest gaps first)
    gaps.sort(key=lambda x: x['gap_score'], reverse=True)
    
    # Calculate weighted overall match
    if gaps:
        total_importance = sum(g['importance'] for g in gaps)
        weighted_match = sum(g['importance'] * g['match_score'] for g in gaps) / total_importance
        weighted_match_pct = weighted_match * 100
    else:
        weighted_match_pct = 100.0
    
    # Identify top missing high-importance skills
    top_missing = [g for g in gaps if not g['matched'] and g['importance'] > 0.7][:3]
    
    return {
        'gaps': gaps,
        'weighted_match_pct': weighted_match_pct,
        'top_missing': top_missing
    }

# ============== EXISTING HELPER FUNCTIONS ==============

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text, len(pdf_reader.pages), "success"
    except Exception as e:
        return "", 0, f"Error: {str(e)}"

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text, len(doc.paragraphs), "success"
    except Exception as e:
        return "", 0, f"Error: {str(e)}"

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    try:
        text = file.read().decode('utf-8')
        return text, len(text.split('\n')), "success"
    except Exception as e:
        return "", 0, f"Error: {str(e)}"

def parse_document(file):
    """Parse uploaded document and extract text"""
    if file is None:
        return None
    
    file_type = file.name.split('.')[-1].lower()
    
    with st.spinner(f'Parsing {file.name}...'):
        if file_type == 'pdf':
            text, pages, status = extract_text_from_pdf(file)
        elif file_type == 'docx':
            text, pages, status = extract_text_from_docx(file)
        elif file_type == 'txt':
            text, pages, status = extract_text_from_txt(file)
        else:
            return None
        
        if status == "success":
            cleaned_text = clean_text(text)
            word_count = len(cleaned_text.split())
            
            return {
                'raw_text': text,
                'cleaned_text': cleaned_text,
                'word_count': word_count,
                'pages': pages,
                'status': status,
                'file_name': file.name,
                'file_type': file_type
            }
        else:
            return {'status': status, 'file_name': file.name}

def clean_text(text):
    """Clean and normalize text"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.,;:!?()\-]', '', text)
    text = re.sub(r'\S+@\S+', '', text)
    text = re.sub(r'\+?\d[\d\-]{7,}\b', '', text)
    text = re.sub(r'http\S+|www\.\S+', '', text)
    return text.strip()

def show_authenticated_sidebar(auth_database):
    """Show sidebar with Profile and Logout buttons"""
    with st.sidebar:
        st.markdown("---")
        
        user_details = auth_database.get_user_details(st.session_state.user['email'])
        
        if user_details and user_details.get('profile_pic') and os.path.exists(user_details.get('profile_pic', '')):
            try:
                profile_img = Image.open(user_details['profile_pic'])
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.image(profile_img, use_container_width=True)
            except:
                st.markdown("### 👤 Welcome!")
        else:
            st.markdown("### 👤 Welcome!")

        st.markdown(
    f"<p style='text-align: center; font-weight: bold;'>{st.session_state.user['name']}</p>",
    unsafe_allow_html=True
)


        if st.button("📝 View/Edit Profile", use_container_width=True):
            st.session_state.show_profile = True
            st.rerun()
        
        st.markdown("---")
        
        # Configuration section
        st.markdown("### ⚙️ Configuration")
        
        # Skill extraction options
        st.markdown("**Skill Extraction**")
        use_noun_chunks = st.checkbox("Use Noun Chunks", value=SPACY_AVAILABLE, 
                                     disabled=not SPACY_AVAILABLE,
                                     help="Extract skills from noun phrases (requires spaCy)")
        
        use_stemming = st.checkbox("Use Stemming", value=False,
                                   help="Normalize words to root form")
        
        confidence_threshold = st.slider(
            "Skill Confidence Threshold",
            0.0, 1.0, 0.6, 0.05,
            help="Minimum confidence for skill extraction"
        )
        
        st.markdown("---")
        st.markdown("**Sentence-BERT**")
        
        enable_sbert = st.checkbox(
            "Enable SBERT Embeddings",
            value=SBERT_AVAILABLE,
            disabled=not SBERT_AVAILABLE,
            help="Use Sentence-BERT for semantic similarity"
        )
        
        if enable_sbert and SBERT_AVAILABLE:
            sbert_model_name = st.selectbox(
                "SBERT Model",
                ['all-MiniLM-L6-v2', 'all-mpnet-base-v2', 'paraphrase-MiniLM-L6-v2'],
                help="Choose embedding model"
            )
            
            normalize_embeddings = st.checkbox("Normalize Embeddings", value=True)
        else:
            sbert_model_name = 'all-MiniLM-L6-v2'
            normalize_embeddings = True
        
        similarity_threshold = st.slider(
            "Similarity Threshold",
            0.0, 1.0, 0.6, 0.05,
            help="Minimum similarity for skill matching"
        )
        
        st.markdown("---")
        
        analyze_btn = st.button("🔍 Analyze", type="primary", use_container_width=True)
        reset_btn = st.button("🔄 Reset", use_container_width=True)
        
        # Logout at bottom
        st.markdown("---")
        if st.button("🚪 Logout", use_container_width=True):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
        
        return {
            'confidence_threshold': confidence_threshold,
            'similarity_threshold': similarity_threshold,
            'analyze_btn': analyze_btn,
            'reset_btn': reset_btn,
            'use_noun_chunks': use_noun_chunks,
            'use_stemming': use_stemming,
            'enable_sbert': enable_sbert,
            'sbert_model_name': sbert_model_name,
            'normalize_embeddings': normalize_embeddings
        }

# ============== ENHANCED VISUALIZATION FUNCTIONS ==============

def create_skill_coverage_gauge(match_percentage):
    """Create a gauge chart for skill coverage"""
    fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = match_percentage,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "Skill Coverage", 'font': {'size': 24}},
        delta = {'reference': 100, 'increasing': {'color': "green"}},
        gauge = {
            'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "darkblue"},
            'bar': {'color': "darkblue"},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': [
                {'range': [0, 50], 'color': 'red'},
                {'range': [50, 75], 'color': 'yellow'},
                {'range': [75, 100], 'color': 'green'}],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': 90}}))
    
    fig.update_layout(height=400)
    return fig

def create_timeline_chart(gap_analysis):
    """Create a timeline for skill development priority"""
    missing_gaps = [g for g in gap_analysis['gaps'] if not g['matched']]
    
    # Sort by importance and gap score
    missing_gaps.sort(key=lambda x: (x['importance'], x['gap_score']), reverse=True)
    
    skills = [g['jd_skill'].title()[:25] for g in missing_gaps[:10]]
    priorities = [g['importance'] * 100 for g in missing_gaps[:10]]
    gap_scores = [g['gap_score'] * 100 for g in missing_gaps[:10]]
    
    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        y=skills,
        x=priorities,
        name='Importance',
        orientation='h',
        marker=dict(color='red', opacity=0.7)
    ))
    
    fig.add_trace(go.Bar(
        y=skills,
        x=gap_scores,
        name='Gap Score',
        orientation='h',
        marker=dict(color='orange', opacity=0.7)
    ))
    
    fig.update_layout(
        title='Skill Development Priority Timeline',
        barmode='overlay',
        height=500,
        xaxis_title='Score',
        yaxis_title='Skills',
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
    )
    
    return fig

def create_sunburst_chart(resume_skills, jd_skills):
    """Create a sunburst chart for skill hierarchy"""
    labels = ['Total Skills']
    parents = ['']
    values = [0]
    colors = ['lightblue']
    
    # Add JD skills
    jd_skill_count = sum(len(skills) for skills in jd_skills.values() if skills)
    labels.append('JD Requirements')
    parents.append('Total Skills')
    values.append(jd_skill_count)
    colors.append('lightcoral')
    
    # Add resume skills
    resume_skill_count = sum(len(skills) for skills in resume_skills.values() if skills)
    labels.append('Resume Skills')
    parents.append('Total Skills')
    values.append(resume_skill_count)
    colors.append('lightgreen')
    
    # Add matched skills
    resume_flat = set()
    for skills in resume_skills.values():
        resume_flat.update(skills.keys())
    
    jd_flat = set()
    for skills in jd_skills.values():
        jd_flat.update(skills.keys())
    
    matched_count = len(resume_flat & jd_flat)
    labels.append('Matched Skills')
    parents.append('JD Requirements')
    values.append(matched_count)
    colors.append('green')
    
    labels.append('Missing Skills')
    parents.append('JD Requirements')
    values.append(jd_skill_count - matched_count)
    colors.append('red')
    
    labels.append('Extra Skills')
    parents.append('Resume Skills')
    values.append(resume_skill_count - matched_count)
    colors.append('blue')
    
    fig = go.Figure(go.Sunburst(
        labels=labels,
        parents=parents,
        values=values,
        marker=dict(colors=colors),
        hovertemplate='<b>%{label}</b><br>Count: %{value}<extra></extra>'
    ))
    
    fig.update_layout(
        title='Skill Hierarchy & Match Overview',
        height=600
    )
    
    return fig

def create_network_graph(gap_analysis, top_n=15):
    """Create a network graph for skill relationships"""
    nodes = []
    edges = []
    
    # Add JD skills as source nodes
    jd_skills = []
    for i, gap in enumerate(gap_analysis['gaps'][:top_n]):
        skill_name = f"JD: {gap['jd_skill'].title()[:20]}"
        nodes.append({
            'id': skill_name,
            'label': gap['jd_skill'].title()[:20],
            'color': 'red' if not gap['matched'] else 'green',
            'value': gap['importance'] * 10,
            'group': 'jd'
        })
        jd_skills.append(skill_name)
    
    # Add matched resume skills as target nodes
    for i, gap in enumerate(gap_analysis['gaps'][:top_n]):
        if gap['matched'] and gap['best_resume_match']:
            resume_skill_name = f"Resume: {gap['best_resume_match'].title()[:20]}"
            if resume_skill_name not in [node['id'] for node in nodes]:
                nodes.append({
                    'id': resume_skill_name,
                    'label': gap['best_resume_match'].title()[:20],
                    'color': 'blue',
                    'value': gap['match_score'] * 10,
                    'group': 'resume'
                })
            
            # Create edge
            jd_skill_name = f"JD: {gap['jd_skill'].title()[:20]}"
            edges.append({
                'from': jd_skill_name,
                'to': resume_skill_name,
                'value': gap['match_score'] * 5,
                'color': f"rgba(0, 255, 0, {gap['match_score']})"
            })
    
    return nodes, edges

def create_progress_bars(gap_analysis):
    """Create progress bars for different skill categories"""
    categories = {}
    
    for gap in gap_analysis['gaps']:
        cat = gap['category'].replace('_', ' ').title()
        if cat not in categories:
            categories[cat] = {'total': 0, 'matched': 0}
        
        categories[cat]['total'] += 1
        if gap['matched']:
            categories[cat]['matched'] += 1
    
    # Calculate percentages
    category_data = []
    for cat, data in categories.items():
        percentage = (data['matched'] / data['total']) * 100 if data['total'] > 0 else 0
        category_data.append({
            'category': cat,
            'matched': data['matched'],
            'total': data['total'],
            'percentage': percentage
        })
    
    # Sort by percentage
    category_data.sort(key=lambda x: x['percentage'])
    
    fig = go.Figure()
    
    for data in category_data:
        fig.add_trace(go.Bar(
            y=[data['category']],
            x=[data['percentage']],
            orientation='h',
            name=data['category'],
            text=f"{data['matched']}/{data['total']} ({data['percentage']:.1f}%)",
            textposition='auto',
            marker=dict(
                color='green' if data['percentage'] > 75 else 
                      'yellow' if data['percentage'] > 50 else 'red'
            )
        ))
    
    fig.update_layout(
        title='Match Percentage by Category',
        xaxis_title='Match Percentage (%)',
        yaxis_title='Categories',
        height=400,
        showlegend=False
    )
    
    return fig

# ============== PDF REPORT GENERATION ==============

def create_pdf_report(gap_analysis, resume_skills, jd_skills, resume_file, jd_file):
    """Create a comprehensive PDF report"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1f77b4'),
            alignment=1
        )
        
        story.append(Paragraph("Resume vs Job Description Analysis Report", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"Resume: {resume_file}", styles['Normal']))
        story.append(Paragraph(f"Job Description: {jd_file}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Page 1: Overall Summary
        story.append(Paragraph("Page 1: Executive Summary", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Overall match metrics
        overall_data = [
            ['Metric', 'Value'],
            ['Weighted Match Score', f"{gap_analysis['weighted_match_pct']:.1f}%"],
            ['Total JD Skills', len(gap_analysis['gaps'])],
            ['Matched Skills', sum(1 for g in gap_analysis['gaps'] if g['matched'])],
            ['Missing Skills', sum(1 for g in gap_analysis['gaps'] if not g['matched'])],
            ['High Priority Gaps', len(gap_analysis['top_missing'])]
        ]
        
        overall_table = Table(overall_data, colWidths=[3*inch, 2*inch])
        overall_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(overall_table)
        story.append(Spacer(1, 20))
        
        # Top 5 gaps
        story.append(Paragraph("Top 5 Skill Gaps:", styles['Heading3']))
        gap_data = [['Skill', 'Category', 'Importance', 'Gap Score']]
        for gap in gap_analysis['top_missing'][:5]:
            gap_data.append([
                gap['jd_skill'].title(),
                gap['category'].replace('_', ' ').title(),
                f"{gap['importance']:.0%}",
                f"{gap['gap_score']:.2f}"
            ])
        
        gap_table = Table(gap_data, colWidths=[1.5*inch, 1.5*inch, 1*inch, 1*inch])
        gap_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightcoral),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.mistyrose),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(gap_table)
        story.append(Spacer(1, 20))
        
        # Recommendations
        story.append(Paragraph("Recommendations:", styles['Heading3']))
        if gap_analysis['weighted_match_pct'] >= 80:
            rec_text = """
            <b>Strong Candidate Match (≥80%)</b><br/>
            • Excellent alignment with job requirements<br/>
            • Proceed to interview stage<br/>
            • Focus on cultural fit and soft skills assessment<br/>
            • Verify experience depth in key areas
            """
        elif gap_analysis['weighted_match_pct'] >= 60:
            rec_text = """
            <b>Good Match with Development Areas (60-80%)</b><br/>
            • Solid foundation with some gaps<br/>
            • Consider for interview with skill assessment<br/>
            • May require onboarding/training for missing skills<br/>
            • Evaluate learning capability and adaptability
            """
        else:
            rec_text = """
            <b>Needs Further Evaluation (<60%)</b><br/>
            • Significant skill gaps identified<br/>
            • Consider junior position or extended training program<br/>
            • Verify if experience compensates for missing skills<br/>
            • Assess candidate's learning curve and motivation
            """
        
        story.append(Paragraph(rec_text, styles['Normal']))
        
        # Add page break
        story.append(Spacer(1, 20))
        
        # Page 2: Detailed Analysis
        story.append(Paragraph("Page 2: Detailed Skill Analysis", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Create a simple heatmap explanation
        heatmap_explanation = """
        <b>Skill Match Heatmap Analysis:</b><br/>
        The heatmap visualization shows the semantic similarity between resume skills and job description requirements. 
        Darker green colors indicate higher similarity matches, while lighter colors show weaker connections. 
        This helps identify both exact matches and related skills that may be transferable.
        """
        story.append(Paragraph(heatmap_explanation, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Skill categories summary
        story.append(Paragraph("Skill Coverage by Category:", styles['Heading3']))
        
        # Calculate category coverage
        categories = {}
        for gap in gap_analysis['gaps']:
            cat = gap['category'].replace('_', ' ').title()
            if cat not in categories:
                categories[cat] = {'total': 0, 'matched': 0}
            categories[cat]['total'] += 1
            if gap['matched']:
                categories[cat]['matched'] += 1
        
        category_data = [['Category', 'Matched', 'Total', 'Coverage']]
        for cat, data in categories.items():
            coverage = (data['matched'] / data['total']) * 100 if data['total'] > 0 else 0
            category_data.append([
                cat,
                str(data['matched']),
                str(data['total']),
                f"{coverage:.1f}%"
            ])
        
        category_table = Table(category_data, colWidths=[2*inch, 1*inch, 1*inch, 1*inch])
        category_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.aliceblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(category_table)
        
        # Build PDF
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
        
    except ImportError:
        # Fallback if reportlab is not available
        return None

# ============== MAIN APPLICATION ==============

def main():
    """Main application entry point"""
    
    init_session_state()
    
    auth_database = AuthDatabase()
    email_service = EmailService()
    
    email_service.configure(
        sender_email="anushachakkiralaa@gmail.com",  
        sender_password="jrjsmueeqdrzwjdr"    
    )
    
    if not st.session_state.authenticated:
        if st.session_state.auth_page == 'login':
            show_login_page(auth_database, email_service)
        elif st.session_state.auth_page == 'signup':
            show_signup_page(auth_database)
        elif st.session_state.auth_page == 'otp':
            show_otp_page(auth_database)
        return
    
    if st.session_state.show_profile:
        show_profile_page(auth_database)
        return
    
    st.set_page_config(
        page_title="SkillGap AI Analyzer",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
        .main-header {
            font-size: 2.5rem;
            font-weight: bold;
            color: #1f77b4;
            text-align: center;
            margin-bottom: 2rem;
        }
        .skill-tag {
            display: inline-block;
            padding: 5px 15px;
            margin: 5px;
            border-radius: 20px;
            background-color: #e3f2fd;
            color: #1976d2;
            font-weight: 500;
        }
        .gap-tag-high {
            display: inline-block;
            padding: 5px 15px;
            margin: 5px;
            border-radius: 20px;
            background-color: #ffebee;
            color: #c62828;
            font-weight: 500;
        }
        .gap-tag-medium {
            display: inline-block;
            padding: 5px 15px;
            margin: 5px;
            border-radius: 20px;
            background-color: #fff3e0;
            color: #e65100;
            font-weight: 500;
        }
        .gap-tag-low {
            display: inline-block;
            padding: 5px 15px;
            margin: 5px;
            border-radius: 20px;
            background-color: #e8f5e9;
            color: #2e7d32;
            font-weight: 500;
        }
    </style>
    """, unsafe_allow_html=True)
    
    sidebar_config = show_authenticated_sidebar(auth_database)
    
    if sidebar_config['reset_btn']:
        st.session_state.parsed_data = {}
        st.session_state.extracted_skills = {}
        st.session_state.analysis_complete = False
        st.session_state.embeddings = {}
        st.session_state.sbert_model = None
        st.rerun()
    
    st.markdown("<div class='main-header'> SKILLGAP AI ANALYZER FOR RESUME AND JOB POST</div>", unsafe_allow_html=True)
    
    tabs = st.tabs([
        "📄 Upload Documents",
        "🔍 Skill Extraction",
        "📊 Skill Gap Analysis",
        "🧠 Sentence-BERT Embeddings",
        "📈 Visualizations",
        "📋 Report"
    ])
    
    # Process files when analyze button is clicked
    if sidebar_config['analyze_btn']:
        with st.spinner("Processing documents..."):
            if 'resume' not in st.session_state.parsed_data or 'jd' not in st.session_state.parsed_data:
                st.warning("Please upload and process both Resume and Job Description first.")
            else:
                # Load SBERT model if enabled
                if sidebar_config['enable_sbert'] and SBERT_AVAILABLE:
                    if st.session_state.sbert_model is None:
                        st.session_state.sbert_model = load_sbert_model(sidebar_config['sbert_model_name'])
                
                # Extract skills
                resume_skills = extract_skills_enhanced(
                    st.session_state.parsed_data['resume']['cleaned_text'],
                    sidebar_config['confidence_threshold'],
                    sidebar_config['use_noun_chunks'],
                    sidebar_config['use_stemming']
                )
                
                jd_skills = extract_skills_enhanced(
                    st.session_state.parsed_data['jd']['cleaned_text'],
                    sidebar_config['confidence_threshold'],
                    sidebar_config['use_noun_chunks'],
                    sidebar_config['use_stemming']
                )
                
                st.session_state.extracted_skills['resume'] = resume_skills
                st.session_state.extracted_skills['jd'] = jd_skills
                
                # Compute embeddings if enabled
                if sidebar_config['enable_sbert'] and st.session_state.sbert_model:
                    # Get all skills for embedding
                    resume_skill_list = []
                    jd_skill_list = []
                    
                    for category, skills in resume_skills.items():
                        resume_skill_list.extend(skills.keys())
                    
                    for category, skills in jd_skills.items():
                        jd_skill_list.extend(skills.keys())
                    
                    if resume_skill_list and jd_skill_list:
                        embeddings_resume = compute_embeddings(
                            resume_skill_list, 
                            st.session_state.sbert_model,
                            sidebar_config['normalize_embeddings']
                        )
                        embeddings_jd = compute_embeddings(
                            jd_skill_list,
                            st.session_state.sbert_model,
                            sidebar_config['normalize_embeddings']
                        )
                        
                        st.session_state.embeddings['resume'] = embeddings_resume
                        st.session_state.embeddings['jd'] = embeddings_jd
                        st.session_state.embeddings['resume_skills'] = resume_skill_list
                        st.session_state.embeddings['jd_skills'] = jd_skill_list
                
                st.session_state.analysis_complete = True
                st.success("✅ Analysis complete!")
    
    # Tab: Upload Documents
    with tabs[0]:
        st.markdown("### Upload Documents")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Upload Resume")
            uploaded_resume = st.file_uploader("Choose resume file (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key='resume_uploader')
            if uploaded_resume:
                st.info(f"Selected: {uploaded_resume.name}")
                if st.button("Process Resume", key="process_resume"):
                    with st.spinner("Processing resume..."):
                        resume_data = parse_document(uploaded_resume)
                        if resume_data:
                            st.session_state.parsed_data['resume'] = resume_data
                            st.success("Resume processed successfully!")
        
        with col2:
            st.markdown("#### Upload Job Description")
            uploaded_jd = st.file_uploader("Choose job description file (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key='jd_uploader')
            if uploaded_jd:
                st.info(f"Selected: {uploaded_jd.name}")
                if st.button("Process Job Description", key="process_jd"):
                    with st.spinner("Processing job description..."):
                        jd_data = parse_document(uploaded_jd)
                        if jd_data:
                            st.session_state.parsed_data['jd'] = jd_data
                            st.success("Job description processed successfully!")
        
        st.markdown("---")
        
        if 'resume' in st.session_state.parsed_data and 'jd' in st.session_state.parsed_data:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### Resume Preview")
                st.info(f"{st.session_state.parsed_data['resume']['file_name']}")
                st.metric("Word Count", st.session_state.parsed_data['resume']['word_count'])
                
                with st.expander("View Text"):
                    st.text_area(
                        "Cleaned Text",
                        st.session_state.parsed_data['resume']['cleaned_text'],
                        height=300,
                        key='resume_preview'
                    )
            
            with col2:
                st.markdown("#### Job Description Preview")
                st.info(f"{st.session_state.parsed_data['jd']['file_name']}")
                st.metric("Word Count", st.session_state.parsed_data['jd']['word_count'])
                
                with st.expander("View Text"):
                    st.text_area(
                        "Cleaned Text",
                        st.session_state.parsed_data['jd']['cleaned_text'],
                        height=300,
                        key='jd_preview'
                    )
    
    # Display results if analysis is complete
    if st.session_state.analysis_complete:
        resume_skills = st.session_state.extracted_skills['resume']
        jd_skills = st.session_state.extracted_skills['jd']
        
        # Tab: Skill Extraction
        with tabs[1]:
            st.markdown("### Extracted Skills by Category")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### Resume Skills")
                
                # Count skills by category
                category_counts_resume = {}
                for category, skills in resume_skills.items():
                    if skills:
                        category_counts_resume[category] = len(skills)
                
                if category_counts_resume:
                    # Show summary
                    st.markdown(f"**Total Categories:** {len(category_counts_resume)}")
                    st.markdown(f"**Total Skills:** {sum(category_counts_resume.values())}")
                    
                    # Show by category
                    for category, skills in resume_skills.items():
                        if skills:
                            category_display = category.replace('_', ' ').title()
                            with st.expander(f"📁 {category_display} ({len(skills)} skills)"):
                                for skill, data in sorted(skills.items(), key=lambda x: x[1]['confidence'], reverse=True):
                                    st.markdown(f"**{skill.title()}**")
                                    st.write(f"- Confidence: {data['confidence']:.2f}")
                                    st.write(f"- Count: {data['count']}")
                                    st.write(f"- Method: {data['extraction_method']}")
                                    if data['evidence']:
                                        st.write(f"- Evidence: _{data['evidence'][0][:100]}..._")
                                    st.markdown("---")
                else:
                    st.info("No skills extracted yet.")
            
            with col2:
                st.markdown("#### Job Description Requirements")
                
                category_counts_jd = {}
                for category, skills in jd_skills.items():
                    if skills:
                        category_counts_jd[category] = len(skills)
                
                if category_counts_jd:
                    st.markdown(f"**Total Categories:** {len(category_counts_jd)}")
                    st.markdown(f"**Total Skills:** {sum(category_counts_jd.values())}")
                    
                    # Flatten resume skills for matching
                    resume_skills_flat = set()
                    for category, skills in resume_skills.items():
                        resume_skills_flat.update(skills.keys())
                    
                    for category, skills in jd_skills.items():
                        if skills:
                            category_display = category.replace('_', ' ').title()
                            with st.expander(f"📁 {category_display} ({len(skills)} skills)"):
                                for skill, data in sorted(skills.items(), key=lambda x: x[1]['confidence'], reverse=True):
                                    match_icon = "✅" if skill in resume_skills_flat else "❌"
                                    st.markdown(f"**{skill.title()}** {match_icon}")
                                    st.write(f"- Confidence: {data['confidence']:.2f}")
                                    st.write(f"- Count: {data['count']}")
                                    st.write(f"- Method: {data['extraction_method']}")
                                    if data['evidence']:
                                        st.write(f"- Evidence: _{data['evidence'][0][:100]}..._")
                                    st.markdown("---")
                else:
                    st.info("No skills extracted yet.")
        
        # Tab: Skill Gap Analysis
        with tabs[2]:
            st.markdown("### Enhanced Skill Gap Analysis")
            
            # Perform gap analysis
            embeddings_resume = st.session_state.embeddings.get('resume')
            embeddings_jd = st.session_state.embeddings.get('jd')
            
            gap_analysis = perform_enhanced_gap_analysis(
                resume_skills,
                jd_skills,
                sidebar_config['similarity_threshold'],
                embeddings_resume,
                embeddings_jd
            )
            
            # Summary Card
            st.markdown("### 📊 Overall Match Summary")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric(
                    "Weighted Match Score",
                    f"{gap_analysis['weighted_match_pct']:.1f}%",
                    help="Overall match considering skill importance"
                )
            
            with col2:
                matched_count = sum(1 for g in gap_analysis['gaps'] if g['matched'])
                total_count = len(gap_analysis['gaps'])
                st.metric(
                    "Matched Skills",
                    f"{matched_count} / {total_count}"
                )
            
            with col3:
                high_importance_gaps = sum(1 for g in gap_analysis['gaps'] if not g['matched'] and g['importance'] > 0.7)
                st.metric(
                    "High-Priority Gaps",
                    high_importance_gaps,
                    delta=f"-{high_importance_gaps}" if high_importance_gaps > 0 else "0"
                )
            
            # Top 3 Missing High-Importance Skills
            if gap_analysis['top_missing']:
                st.markdown("### 🔴 Top 3 Missing High-Importance Skills")
                cols = st.columns(3)
                for idx, gap in enumerate(gap_analysis['top_missing']):
                    with cols[idx]:
                        st.markdown(f"""
                        <div style='background:#ffebee; padding:15px; border-radius:10px; border-left:4px solid #c62828;'>
                            <h4 style='margin:0; color:#c62828;'>{gap['jd_skill'].title()}</h4>
                            <p style='margin:5px 0;'><b>Category:</b> {gap['category'].replace('_', ' ').title()}</p>
                            <p style='margin:5px 0;'><b>Importance:</b> {gap['importance']:.0%}</p>
                            <p style='margin:5px 0;'><b>Gap Score:</b> {gap['gap_score']:.2f}</p>
                        </div>
                        """, unsafe_allow_html=True)
            
            st.markdown("---")
            
            # Detailed Gap Table
            st.markdown("### 📋 Detailed Skill Gap Analysis")
            
            # Create detailed dataframe
            gap_df_data = []
            for gap in gap_analysis['gaps']:
                gap_df_data.append({
                    'JD Skill': gap['jd_skill'].title(),
                    'Category': gap['category'].replace('_', ' ').title(),
                    'Type': gap['type'].title(),
                    'Importance': f"{gap['importance']:.0%}",
                    'Match Score': f"{gap['match_score']:.2f}",
                    'Best Resume Match': gap['best_resume_match'].title() if gap['best_resume_match'] else 'None',
                    'Status': '✅ Matched' if gap['matched'] else '❌ Missing',
                    'Gap Score': f"{gap['gap_score']:.2f}"
                })
            
            gap_df = pd.DataFrame(gap_df_data)
            
            # Add filters
            col1, col2, col3 = st.columns(3)
            with col1:
                filter_status = st.multiselect(
                    "Filter by Status",
                    ['✅ Matched', '❌ Missing'],
                    default=['❌ Missing']
                )
            with col2:
                filter_type = st.multiselect(
                    "Filter by Type",
                    gap_df['Type'].unique(),
                    default=gap_df['Type'].unique()
                )
            with col3:
                sort_by = st.selectbox(
                    "Sort by",
                    ['Gap Score', 'Importance', 'Match Score'],
                    index=0
                )
            
            # Apply filters
            filtered_df = gap_df[
                (gap_df['Status'].isin(filter_status)) &
                (gap_df['Type'].isin(filter_type))
            ]
            
            # Sort
            sort_ascending = False if sort_by == 'Gap Score' else False
            filtered_df = filtered_df.sort_values(by=sort_by, ascending=sort_ascending)
            
            st.dataframe(filtered_df, use_container_width=True, height=400)
        
        # Tab: Sentence-BERT Embeddings
        with tabs[3]:
            st.markdown("### 🧠 Sentence-BERT Embeddings Analysis")
            
            if not SBERT_AVAILABLE:
                st.error("⚠️ Sentence-BERT is not available. Install with: `pip install sentence-transformers`")
            elif not sidebar_config['enable_sbert']:
                st.info("ℹ️ Enable SBERT in the sidebar to use embedding-based analysis")
            elif 'resume' not in st.session_state.embeddings or 'jd' not in st.session_state.embeddings:
                st.warning("⚠️ Run analysis first to compute embeddings")
            else:
                st.success(f"✅ Using model: {sidebar_config['sbert_model_name']}")
                
                embeddings_resume = st.session_state.embeddings['resume']
                embeddings_jd = st.session_state.embeddings['jd']
                resume_skill_list = st.session_state.embeddings['resume_skills']
                jd_skill_list = st.session_state.embeddings['jd_skills']
                
                # Compute similarity matrix
                similarity_matrix = compute_similarity_matrix(embeddings_jd, embeddings_resume)
                
                if similarity_matrix is not None:
                    # Show embedding statistics
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Resume Embeddings", len(resume_skill_list))
                    with col2:
                        st.metric("JD Embeddings", len(jd_skill_list))
                    with col3:
                        avg_similarity = similarity_matrix.mean()
                        st.metric("Avg Similarity", f"{avg_similarity:.3f}")
                    
                    st.markdown("---")
                    
                    # Similarity heatmap
                    st.markdown("### Semantic Similarity Heatmap")
                    
                    # Limit to top 20 skills for visualization
                    display_limit = 20
                    jd_display = jd_skill_list[:display_limit]
                    resume_display = resume_skill_list[:display_limit]
                    sim_matrix_display = similarity_matrix[:len(jd_display), :len(resume_display)]
                    
                    fig = go.Figure(data=go.Heatmap(
                        z=sim_matrix_display,
                        x=[s[:20] for s in resume_display],
                        y=[s[:20] for s in jd_display],
                        colorscale='RdYlGn',
                        zmid=sidebar_config['similarity_threshold']
                    ))
                    
                    fig.update_layout(
                        title=f'Semantic Similarity Matrix (Top {display_limit} skills)',
                        xaxis_title='Resume Skills',
                        yaxis_title='JD Skills',
                        height=600
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    st.markdown("---")
                    
                    # Best matches
                    st.markdown("### 🎯 Best Semantic Matches")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### Top Resume-JD Matches")
                        matches = []
                        for i, jd_skill in enumerate(jd_skill_list):
                            if i < len(similarity_matrix):
                                best_idx = np.argmax(similarity_matrix[i])
                                best_score = similarity_matrix[i][best_idx]
                                if best_idx < len(resume_skill_list):
                                    matches.append((jd_skill, resume_skill_list[best_idx], best_score))
                        
                        matches.sort(key=lambda x: x[2], reverse=True)
                        
                        for jd_s, res_s, score in matches[:10]:
                            color = "#4caf50" if score > 0.8 else "#ff9800" if score > 0.6 else "#f44336"
                            st.markdown(f"""
                            <div style='padding:10px; margin:5px 0; border-left:4px solid {color}; background:#f5f5f5;'>
                                <b>JD:</b> {jd_s.title()}<br>
                                <b>Resume:</b> {res_s.title()}<br>
                                <b>Similarity:</b> {score:.3f}
                            </div>
                            """, unsafe_allow_html=True)
                    
                    with col2:
                        st.markdown("#### Similarity Distribution")
                        
                        all_similarities = similarity_matrix.flatten()
                        
                        fig = px.histogram(
                            all_similarities,
                            nbins=50,
                            title='Distribution of Similarity Scores',
                            labels={'value': 'Similarity Score', 'count': 'Frequency'},
                            color_discrete_sequence=['#1f77b4']
                        )
                        
                        fig.add_vline(
                            x=sidebar_config['similarity_threshold'],
                            line_dash="dash",
                            line_color="red",
                            annotation_text="Threshold"
                        )
                        
                        st.plotly_chart(fig, use_container_width=True)
        
        # Tab: Visualizations
        with tabs[4]:
            st.markdown("### 📊 Interactive Visualizations")
            
            # Flatten skills for counting
            resume_skills_flat = {}
            jd_skills_flat = {}
            
            for category, skills in resume_skills.items():
                for skill, data in skills.items():
                    resume_skills_flat[skill] = data
            
            for category, skills in jd_skills.items():
                for skill, data in skills.items():
                    jd_skills_flat[skill] = data
            
            # Perform gap analysis for new visualizations
            gap_analysis = perform_enhanced_gap_analysis(
                resume_skills,
                jd_skills,
                sidebar_config['similarity_threshold'],
                st.session_state.embeddings.get('resume'),
                st.session_state.embeddings.get('jd')
            )
            
            # New visualization tabs
            viz_tabs = st.tabs([
                "📈 Overview",
                "🎯 Gap Analysis", 
                "🌐 Network View",
                "📊 Category Analysis",
                "📅 Development Plan"
            ])
            
            with viz_tabs[0]:
                col1, col2 = st.columns(2)
                
                with col1:
                    # Gauge chart
                    st.plotly_chart(
                        create_skill_coverage_gauge(gap_analysis['weighted_match_pct']), 
                        use_container_width=True
                    )
                    
                    # Sunburst chart
                    st.plotly_chart(
                        create_sunburst_chart(resume_skills, jd_skills),
                        use_container_width=True
                    )
                
                with col2:
                    # Matched vs Missing
                    matched = len(set(resume_skills_flat.keys()) & set(jd_skills_flat.keys()))
                    missing = len(jd_skills_flat) - matched
                    
                    fig = px.pie(
                        values=[matched, missing],
                        names=['Matched', 'Missing'],
                        title='Overall Match Status',
                        color_discrete_sequence=['#4caf50', '#f44336']
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Skills by category
                    category_counts = {}
                    for category, skills in resume_skills.items():
                        if skills:
                            cat_display = category.replace('_', ' ').title()
                            category_counts[cat_display] = len(skills)
                    
                    if category_counts:
                        fig = px.bar(
                            x=list(category_counts.values()),
                            y=list(category_counts.keys()),
                            orientation='h',
                            title='Resume Skills by Category',
                            labels={'x': 'Count', 'y': 'Category'},
                            color=list(category_counts.values()),
                            color_continuous_scale='Blues'
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
            with viz_tabs[1]:
                col1, col2 = st.columns(2)
                
                with col1:
                    # Progress bars by category
                    st.plotly_chart(
                        create_progress_bars(gap_analysis),
                        use_container_width=True
                    )
                    
                    # Top gaps by gap score
                    top_gaps = [g for g in gap_analysis['gaps'] if not g['matched']][:15]
                    if top_gaps:
                        gap_skills = [g['jd_skill'].title()[:30] for g in top_gaps]
                        gap_scores = [g['gap_score'] for g in top_gaps]
                        gap_importance = [g['importance'] for g in top_gaps]
                        
                        fig = go.Figure()
                        fig.add_trace(go.Bar(
                            y=gap_skills,
                            x=gap_scores,
                            orientation='h',
                            marker=dict(
                                color=gap_importance,
                                colorscale='Reds',
                                showscale=True,
                                colorbar=dict(title="Importance")
                            ),
                            text=[f"{score:.2f}" for score in gap_scores],
                            textposition='outside',
                            hovertemplate='<b>%{y}</b><br>Gap Score: %{x:.3f}<extra></extra>'
                        ))
                        
                        fig.update_layout(
                            title='Top 15 Priority Skill Gaps',
                            xaxis_title='Gap Score (Importance × [1 - Match])',
                            yaxis_title='Skill',
                            height=600,
                            yaxis={'categoryorder': 'total ascending'}
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    # Timeline chart
                    st.plotly_chart(
                        create_timeline_chart(gap_analysis),
                        use_container_width=True
                    )
                    
                    # Waterfall chart
                    top_gaps_waterfall = [g for g in gap_analysis['gaps'] if not g['matched']][:10]
                    if top_gaps_waterfall:
                        skills = ['Start'] + [g['jd_skill'].title()[:20] for g in top_gaps_waterfall] + ['Total Gap']
                        measures = ['relative'] + ['relative'] * len(top_gaps_waterfall) + ['total']
                        values = [0] + [g['gap_score'] for g in top_gaps_waterfall] + [sum(g['gap_score'] for g in top_gaps_waterfall)]
                        
                        fig = go.Figure(go.Waterfall(
                            x=skills,
                            y=values,
                            measure=measures,
                            increasing={"marker": {"color": "#f44336"}},
                            totals={"marker": {"color": "#1976d2"}},
                            connector={"line": {"color": "rgb(63, 63, 63)"}},
                            text=[f"{v:.2f}" if v != 0 else "" for v in values],
                            textposition="outside"
                        ))
                        
                        fig.update_layout(
                            title='Top 10 Gap Contributors (Waterfall)',
                            height=500,
                            xaxis_title='Skills',
                            yaxis_title='Cumulative Gap Score'
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
            with viz_tabs[2]:
                st.markdown("### 🔗 Skill Relationship Network")
                
                nodes, edges = create_network_graph(gap_analysis)
                
                if nodes:
                    # Convert to vis.js format
                    vis_nodes = []
                    for node in nodes:
                        vis_nodes.append({
                            'id': node['id'],
                            'label': node['label'],
                            'color': node['color'],
                            'value': node['value']
                        })
                    
                    vis_edges = []
                    for edge in edges:
                        vis_edges.append({
                            'from': edge['from'],
                            'to': edge['to'],
                            'value': edge['value'],
                            'color': edge['color']
                        })
                    
                    # Create HTML for network graph
                    network_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <script type="text/javascript" src="https://unpkg.com/vis-network/standalone/umd/vis-network.min.js"></script>
                        <style>
                            #network {{
                                width: 100%;
                                height: 600px;
                                border: 1px solid lightgray;
                            }}
                        </style>
                    </head>
                    <body>
                        <div id="network"></div>
                        <script type="text/javascript">
                            var nodes = new vis.DataSet({vis_nodes});
                            var edges = new vis.DataSet({vis_edges});
                            
                            var container = document.getElementById('network');
                            var data = {{
                                nodes: nodes,
                                edges: edges
                            }};
                            var options = {{
                                nodes: {{
                                    shape: 'dot',
                                    size: 25,
                                    font: {{
                                        size: 14,
                                        color: '#000000'
                                    }},
                                    borderWidth: 2
                                }},
                                edges: {{
                                    width: 2,
                                    smooth: true
                                }},
                                physics: {{
                                    enabled: true,
                                    stabilization: true
                                }},
                                interaction: {{
                                    hover: true,
                                    tooltipDelay: 200
                                }}
                            }};
                            
                            var network = new vis.Network(container, data, options);
                        </script>
                    </body>
                    </html>
                    """
                    
                    components.html(network_html, height=600)
                    
                    st.markdown("""
                    **Network Legend:**
                    - 🔴 **Red Nodes**: Missing JD skills
                    - 🟢 **Green Nodes**: Matched JD skills  
                    - 🔵 **Blue Nodes**: Resume skills
                    - 🟢 **Green Edges**: Match connections (darker = better match)
                    """)
                else:
                    st.info("No significant skill relationships to display.")
            
            with viz_tabs[3]:
                # Radar chart comparison
                st.markdown("### Category-wise Comparison")
                
                categories = []
                resume_counts = []
                jd_counts = []
                
                for category in CATEGORIZED_SKILLS.keys():
                    cat_display = category.replace('_', ' ').title()
                    categories.append(cat_display)
                    resume_counts.append(len(resume_skills.get(category, {})))
                    jd_counts.append(len(jd_skills.get(category, {})))
                
                fig = go.Figure()
                
                fig.add_trace(go.Scatterpolar(
                    r=resume_counts,
                    theta=categories,
                    fill='toself',
                    name='Resume',
                    line_color='#1f77b4'
                ))
                
                fig.add_trace(go.Scatterpolar(
                    r=jd_counts,
                    theta=categories,
                    fill='toself',
                    name='Job Description',
                    line_color='#ff7f0e'
                ))
                
                fig.update_layout(
                    polar=dict(radialaxis=dict(visible=True, range=[0, max(max(resume_counts), max(jd_counts))+2])),
                    showlegend=True,
                    height=600,
                    title='Skills Distribution Across Categories'
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                # Additional category analysis
                col1, col2 = st.columns(2)
                
                with col1:
                    # Missing skills by category
                    category_gaps = {}
                    for gap in gap_analysis['gaps']:
                        if not gap['matched']:
                            cat = gap['category'].replace('_', ' ').title()
                            category_gaps[cat] = category_gaps.get(cat, 0) + 1
                    
                    if category_gaps:
                        # Sort by count
                        sorted_cats = sorted(category_gaps.items(), key=lambda x: x[1], reverse=True)
                        cats = [x[0] for x in sorted_cats]
                        counts = [x[1] for x in sorted_cats]
                        
                        fig = px.bar(
                            x=counts,
                            y=cats,
                            orientation='h',
                            title='Missing Skills by Category',
                            labels={'x': 'Number of Missing Skills', 'y': 'Category'},
                            color=counts,
                            color_continuous_scale='Reds',
                            text=counts
                        )
                        fig.update_traces(textposition='outside')
                        fig.update_layout(height=500, showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    # Pie chart: Gap severity distribution
                    if gap_analysis['gaps']:
                        high_gaps = sum(1 for g in gap_analysis['gaps'] if not g['matched'] and g['importance'] > 0.7)
                        medium_gaps = sum(1 for g in gap_analysis['gaps'] if not g['matched'] and 0.5 <= g['importance'] <= 0.7)
                        low_gaps = sum(1 for g in gap_analysis['gaps'] if not g['matched'] and g['importance'] < 0.5)
                        matched_count = sum(1 for g in gap_analysis['gaps'] if g['matched'])
                        
                        fig = go.Figure(data=[go.Pie(
                            labels=['Matched', 'High Priority Gap', 'Medium Priority Gap', 'Low Priority Gap'],
                            values=[matched_count, high_gaps, medium_gaps, low_gaps],
                            marker=dict(colors=['#4caf50', '#d32f2f', '#ff9800', '#ffc107']),
                            hole=0.4
                        )])
                        
                        fig.update_layout(
                            title='Gap Severity Distribution',
                            height=500,
                            annotations=[dict(text=f'{len(gap_analysis["gaps"])} Total', x=0.5, y=0.5, 
                                            font_size=20, showarrow=False)]
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
            with viz_tabs[4]:
                st.markdown("### 📅 Skill Development Roadmap")
                
                # Development priority matrix
                missing_gaps = [g for g in gap_analysis['gaps'] if not g['matched']]
                
                if missing_gaps:
                    # Sort by importance and gap score
                    missing_gaps.sort(key=lambda x: (x['importance'], x['gap_score']), reverse=True)
                    
                    # Create development phases
                    phases = {
                        'Immediate (1-3 months)': [g for g in missing_gaps if g['importance'] > 0.8][:3],
                        'Short-term (3-6 months)': [g for g in missing_gaps if 0.6 <= g['importance'] <= 0.8][:4],
                        'Medium-term (6-12 months)': [g for g in missing_gaps if 0.4 <= g['importance'] < 0.6][:5],
                        'Long-term (12+ months)': [g for g in missing_gaps if g['importance'] < 0.4][:3]
                    }
                    
                    for phase, skills in phases.items():
                        if skills:
                            st.markdown(f"#### {phase}")
                            
                            for skill in skills:
                                col1, col2, col3 = st.columns([3, 1, 1])
                                with col1:
                                    st.write(f"**{skill['jd_skill'].title()}**")
                                with col2:
                                    st.write(f"Priority: {skill['importance']:.0%}")
                                with col3:
                                    st.write(f"Gap: {skill['gap_score']:.2f}")
                            
                            st.markdown("---")
                    
                    # Learning path visualization
                    st.markdown("#### 📚 Recommended Learning Path")
                    
                    learning_data = []
                    for phase, skills in phases.items():
                        for skill in skills:
                            learning_data.append({
                                'Skill': skill['jd_skill'].title(),
                                'Timeline': phase,
                                'Priority': skill['importance'],
                                'Gap Score': skill['gap_score']
                            })
                    
                    if learning_data:
                        learning_df = pd.DataFrame(learning_data)
                        
                        fig = px.timeline(
                            learning_df,
                            x_start=[0]*len(learning_df),
                            x_end=learning_df['Priority'] * 100,
                            y=learning_df['Timeline'],
                            color=learning_df['Gap Score'],
                            hover_data=['Skill', 'Priority'],
                            title='Skill Development Timeline',
                            color_continuous_scale='Reds'
                        )
                        
                        fig.update_layout(height=400)
                        st.plotly_chart(fig, use_container_width=True)
                else:
                    st.success("🎉 No skill gaps identified! All required skills are covered.")
        
        # Tab: Report
        with tabs[5]:
            st.markdown("### 📋 Comprehensive Match Report")
            
            report_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # Download buttons
            col1, col2, col3 = st.columns(3)
            
            gap_analysis = perform_enhanced_gap_analysis(
                resume_skills, jd_skills, sidebar_config['similarity_threshold'],
                st.session_state.embeddings.get('resume'),
                st.session_state.embeddings.get('jd')
            )
            
            # Prepare export data
            export_data = []
            for gap in gap_analysis['gaps']:
                export_data.append({
                    'JD_Skill': gap['jd_skill'],
                    'Category': gap['category'],
                    'Type': gap['type'],
                    'Importance': gap['importance'],
                    'Match_Score': gap['match_score'],
                    'Best_Resume_Match': gap['best_resume_match'] if gap['best_resume_match'] else 'None',
                    'Status': 'Matched' if gap['matched'] else 'Missing',
                    'Gap_Score': gap['gap_score']
                })
            
            df_export = pd.DataFrame(export_data)
            csv = df_export.to_csv(index=False)
            
            with col1:
                st.download_button(
                    label='📥 Download Detailed CSV Report',
                    data=csv,
                    file_name=f'skill_gap_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv',
                    mime='text/csv',
                    use_container_width=True
                )
            
            with col2:
                report_text = f"""RESUME-JD MATCH REPORT
Generated: {report_date}

OVERALL MATCH SUMMARY
====================
Weighted Match Score: {gap_analysis['weighted_match_pct']:.1f}%
Total JD Skills: {len(gap_analysis['gaps'])}
Matched Skills: {sum(1 for g in gap_analysis['gaps'] if g['matched'])}
Missing Skills: {sum(1 for g in gap_analysis['gaps'] if not g['matched'])}

TOP MISSING HIGH-IMPORTANCE SKILLS
==================================
{chr(10).join([f"{i+1}. {g['jd_skill'].title()} (Importance: {g['importance']:.0%}, Gap Score: {g['gap_score']:.2f})" for i, g in enumerate(gap_analysis['top_missing'])])}

DETAILED SKILL ANALYSIS
========================
{chr(10).join([f"- {g['jd_skill'].title()}: {'✓ Matched' if g['matched'] else '✗ Missing'} (Score: {g['match_score']:.2f})" for g in gap_analysis['gaps'][:20]])}
"""
                
                st.download_button(
                    label='📥 Download Text Report',
                    data=report_text,
                    file_name=f'match_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt',
                    mime='text/plain',
                    use_container_width=True
                )
            
            with col3:
                # PDF Report button
                if st.button('📄 Generate PDF Report', use_container_width=True):
                    with st.spinner('Generating PDF report...'):
                        try:
                            pdf_data = create_pdf_report(
                                gap_analysis,
                                resume_skills,
                                jd_skills,
                                st.session_state.parsed_data['resume']['file_name'],
                                st.session_state.parsed_data['jd']['file_name']
                            )
                            
                            if pdf_data:
                                st.download_button(
                                    label='📥 Download PDF Report',
                                    data=pdf_data,
                                    file_name=f'resume_jd_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                                    mime='application/pdf',
                                    use_container_width=True
                                )
                            else:
                                st.error("PDF generation failed. Please install reportlab: `pip install reportlab`")
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")
                            st.info("PDF generation requires additional dependencies: `pip install reportlab`")
            
            st.markdown("---")
            
            # Display report
            st.markdown(f"""
            ## Resume-Job Description Match Report
            **Generated:** {report_date}
            **Model:** {sidebar_config['sbert_model_name'] if sidebar_config['enable_sbert'] else 'TF-IDF'}
            
            ---
            
            ### 📊 Match Summary
            - **Weighted Match Score:** {gap_analysis['weighted_match_pct']:.1f}%
            - **Matched Skills:** {sum(1 for g in gap_analysis['gaps'] if g['matched'])} / {len(gap_analysis['gaps'])}
            - **High-Priority Gaps:** {len(gap_analysis['top_missing'])}
            
            ---
            
            ### 🎯 Recommendation
            """)
            
            if gap_analysis['weighted_match_pct'] >= 80:
                st.success("""
                **Strong Candidate Match (≥80%)**
                - Excellent alignment with job requirements
                - Proceed to interview stage
                - Focus on cultural fit and soft skills assessment
                """)
            elif gap_analysis['weighted_match_pct'] >= 60:
                st.warning("""
                **Good Match with Development Areas (60-80%)**
                - Solid foundation with some gaps
                - Consider for interview with skill assessment
                - May require onboarding/training for missing skills
                """)
            else:
                st.error("""
                **Needs Further Evaluation (<60%)**
                - Significant skill gaps identified
                - Consider junior position or extended training program
                - Verify if experience compensates for missing skills
                """)
    
    else:
        st.info("""
        ### 👋 Welcome to SkillGap AI Analyzer!
        
        **New Features:**
        - 🔍 **Advanced Skill Extraction**: Noun chunks, NER, and pattern matching
        - 🧠 **Sentence-BERT**: Semantic similarity analysis
        - 📊 **Weighted Gap Analysis**: Importance-based ranking
        - 🎯 **Smart Matching**: Best-match identification
        
        **Get Started:**
        1. Upload Resume and Job Description
        2. Configure extraction and similarity settings in sidebar
        3. Click "Analyze" to start
        4. Review detailed results across tabs
        
        **Required Libraries:**
        - `sentence-transformers` for SBERT
        - `spacy` for NLP features
        - `nltk` for text normalization
        """)
    
    st.markdown("---")

if __name__ == "__main__":
    main()